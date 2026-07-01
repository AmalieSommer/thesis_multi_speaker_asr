import pandas as pd
import os
from docx import Document
import re
import json


def read_file(path: str):
    try:
        doc = Document(path)
    except Exception as e:
        print(f'An error occurred: {e}')
    finally:
        return doc

def extractNumberFiles(doc: Document):
    filesMap = {}

    for i, paragraph in enumerate(doc.paragraphs):
        filenames = paragraph.text.split(':')[1]

        if '+' in filenames:
            filename = filenames.split('+')[0].strip()
            base = filename[:-1] if filename[-1].isdigit() else filename
            num_files = len(filenames.split('+'))

            for i in range(num_files):
                id = i + 1
                filesMap[id] = base + str(id)
        else:
            lst = [filenames]
            id = len(lst)
            filesMap[id] = filenames

        return filesMap

def get_seconds(timestamp: str):
    m, s = timestamp.split(':')
    return float(m) * 60.0 + float(s)  


def generateSegments(doc: Document, filesMap: dict):
    jsonFile = []
    json_line = {}
    transcript_segments = []
    timestamp_pattern = r'\d\d:\d\d' # TODO: Change later for real timestamp pattern...

    flag_ignore_paragraphs = False
    for key, value in filesMap.items():
        json_line['id'] = value # To save the file id of the transcription being passed.

        for p in doc.paragraphs:
            
            match = re.match(r'Fil\s+(\d+)\s+begynder', p.text)
            transcript = {}
                
            if flag_ignore_paragraphs:
                # Should ignore paragraphs until reset...
                if match:
                    if str(key) in match.group():
                        flag_ignore_paragraphs = False
                        continue
            else:
                # Should include paragraphs as normal:

                if match:
                    # Check if the matched line contains the number of the file currently viewed,
                    # otherwise set the ignoring flag to True
                    if str(key) not in match.group():
                        flag_ignore_paragraphs = True
                        continue # ensure the flag is set correctly to not ignore and move onto next paragraph
        
                match_time = re.search('^' + timestamp_pattern, p.text)
                text = p.text
                if match_time:
                    time = match_time.group()
                    transcript['start'] = get_seconds(time)
                    text = re.sub(time, '', text)

                    match_speaker = re.search('(?<=.)[BIV2]:', p.text)
                    if match_speaker:
                        speaker_group = match_speaker.group()
                        transcript['speaker'] = speaker_group if type(speaker_group) == str else speaker_group[0]   # to not include the semicolon...
                        
                        text = re.sub(speaker_group, '', text)
                        transcript['text'] = text

                if len(transcript) != 0:
                    transcript_segments.append(transcript)

        json_line['segments'] = transcript_segments
        
        jsonFile.append(json_line)
        json_line = {}
    return jsonFile

def save(result, filename):
    try:
        with open(os.path.join(os.getcwd(), 'data', f'{filename}.jsonl'), "w") as file:
            for item in result:
                json_line = json.dumps(item)
                file.write(json_line + '\n')
        print(f"Successfully updated the file")
    except Exception as e:
        print(f"Error updating the file: {e}")


if __name__=='__main__':
    print('...')
    filename = 'transcripts'    # To hold the complete transcripts from all recordings

    files_list = os.listdir('data')
    
    # Iterate the list of docx files in the data folder:
    for document in files_list:
        if '.docx' in document:
            path = os.path.join(os.getcwd(), 'data', document)
            doc = read_file(path=path)
            map = extractNumberFiles(doc)
            results = generateSegments(doc, map)
            save(result=results, filename=filename)
    